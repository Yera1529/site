"""Local file storage abstraction with text extraction. Replace this module to support S3/MinIO."""

import os
import uuid
import pdfplumber
import docx2txt
from pathlib import Path
from config import get_settings


class StorageService:
    # Защита от «zip-бомбы»: docx/odt — это zip-архивы. Ограничиваем суммарный
    # распакованный размер и число элементов, иначе вредоносный файл может
    # исчерпать память контейнера при извлечении текста.
    _MAX_UNCOMPRESSED_BYTES = 300 * 1024 * 1024  # 300 МБ суммарно
    _MAX_ZIP_ENTRIES = 5000

    def __init__(self):
        settings = get_settings()
        self.base_dir = Path(settings.storage_dir)
        self.uploads_dir = self.base_dir / "uploads"
        self.generated_dir = self.base_dir / "generated"
        self.uploads_dir.mkdir(parents=True, exist_ok=True)
        self.generated_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """Защита от path traversal: только базовое имя, без разделителей путей
        и спецсимволов. Возвращает безопасное имя (без каталогов)."""
        base = os.path.basename(filename or "")
        base = "".join(c if (c.isalnum() or c in " .-_()№") else "_" for c in base)
        base = base.strip(" .") or "file"      # не допускаем пустое/только-точки имя
        return base[:120]

    @staticmethod
    def _assert_zip_safe(path: str) -> None:
        """Проверяет zip-контейнер (docx/odt) на «zip-бомбу» ДО распаковки.

        Бросает ValueError, если суммарный распакованный размер или число
        элементов превышают лимиты. Вызывающий код перехватывает исключение
        и возвращает безопасное сообщение об ошибке.
        """
        import zipfile
        with zipfile.ZipFile(path, "r") as z:
            infos = z.infolist()
            if len(infos) > StorageService._MAX_ZIP_ENTRIES:
                raise ValueError("Документ содержит слишком много элементов (возможна zip-бомба).")
            total = sum(i.file_size for i in infos)
            if total > StorageService._MAX_UNCOMPRESSED_BYTES:
                raise ValueError("Распакованный размер документа превышает лимит (возможна zip-бомба).")

    def save_file(self, matter_id: str, filename: str, content: bytes) -> str:
        """Save an uploaded file and return its relative storage path."""
        matter_dir = self.uploads_dir / matter_id
        matter_dir.mkdir(parents=True, exist_ok=True)

        safe_name = f"{uuid.uuid4().hex}_{self._sanitize_filename(filename)}"
        file_path = matter_dir / safe_name
        file_path.write_bytes(content)
        return str(file_path.relative_to(self.base_dir))

    def get_full_path(self, relative_path: str) -> str:
        """Resolve a relative storage path to an absolute path."""
        return str(self.base_dir / relative_path)

    def delete_file(self, relative_path: str) -> None:
        """Delete a file from storage."""
        full_path = self.base_dir / relative_path
        if full_path.exists():
            full_path.unlink()

    def extract_text(self, relative_path: str, file_type: str) -> str:
        """Extract text content from an uploaded file."""
        full_path = self.base_dir / relative_path
        try:
            if file_type == "pdf":
                return self._extract_pdf(str(full_path))
            elif file_type == "docx":
                return self._extract_docx(str(full_path))
            elif file_type == "doc":
                return self._extract_doc(str(full_path))
            elif file_type == "txt":
                return full_path.read_text(encoding="utf-8", errors="ignore")
            elif file_type == "rtf":
                return self._extract_rtf(str(full_path))
            elif file_type == "odt":
                return self._extract_odt(str(full_path))
        except Exception as e:
            return f"[Ошибка извлечения текста: {e}]"
        return ""

    @staticmethod
    def _extract_pdf(path: str) -> str:
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(path: str) -> str:
        """Extract text from DOCX using python-docx for full coverage:
        paragraphs, tables (with cell content), and headers/footers.
        docx2txt misses tables which often contain names, dates, case numbers.
        """
        StorageService._assert_zip_safe(path)
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        parts = []

        # Main body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables — critical for extracting FIO, dates, case numbers
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))

        # Headers and footers (may contain investigator name, department)
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                text = hdr_para.text.strip()
                if text:
                    parts.append(f"[Колонтитул: {text}]")
            for ftr_para in section.footer.paragraphs:
                text = ftr_para.text.strip()
                if text:
                    parts.append(f"[Нижний колонтитул: {text}]")

        return "\n".join(parts)

    @staticmethod
    def _extract_doc(path: str) -> str:
        """Extract text from old-format .doc files using docx2txt.
        python-docx does NOT support .doc, only .docx.
        docx2txt handles both formats.
        """
        import zipfile
        # Если .doc на деле OOXML-zip (docx2txt читает его как zip) — проверяем на
        # zip-бомбу. Настоящий бинарный .doc не zip → проверку пропускаем.
        if zipfile.is_zipfile(path):
            StorageService._assert_zip_safe(path)
        try:
            text = docx2txt.process(path)
            if text and text.strip():
                return text.strip()
        except Exception:
            pass
        # Fallback: try reading as plain text (some .doc files are actually RTF)
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read()
            if raw.startswith("{\\rtf"):
                import re
                text = re.sub(r"\\[a-z]+\d*\s?", " ", raw)
                text = re.sub(r"[{}]", "", text)
                text = re.sub(r"\s+", " ", text).strip()
                return text
        except Exception:
            pass
        return "[Не удалось извлечь текст из .doc файла]"


    @staticmethod
    def _extract_rtf(path: str) -> str:
        """Best-effort RTF extraction by stripping RTF control words."""
        import re
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        text = re.sub(r"\\[a-z]+\d*\s?", " ", raw)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    @staticmethod
    def _extract_odt(path: str) -> str:
        """Extract text from ODT (ZIP of XML) files."""
        import zipfile
        import re
        StorageService._assert_zip_safe(path)
        with zipfile.ZipFile(path, "r") as z:
            if "content.xml" not in z.namelist():
                return ""
            xml = z.read("content.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        return text
