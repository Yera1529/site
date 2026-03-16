"""Document DOCX export service with GOST-style formatting and template parsing."""

import re
import uuid
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import get_settings


ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


# ---------------------------------------------------------------------------
# Single-pass HTML→items parser (preserves document order)
# ---------------------------------------------------------------------------

class _DocParser(HTMLParser):
    """Parse HTML into a flat ordered list of document items."""

    BLOCK_TAGS = {"p", "h1", "h2", "h3", "li", "div", "ol", "ul", "br"}

    def __init__(self):
        super().__init__()
        self.items: list[dict] = []
        self._stack: list[dict] = []   # current open block tag stack
        self._buf: list[str] = []      # text buffer for current block
        self._attrs_stack: list[dict] = []  # attributes for each open block

    # ── helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def _decode(text: str) -> str:
        text = text.replace("&nbsp;", " ")
        text = text.replace("&amp;", "&")
        text = text.replace("&lt;", "<")
        text = text.replace("&gt;", ">")
        text = text.replace("&#160;", " ")
        text = re.sub(r"&[a-z]+;", "", text)
        return text

    @staticmethod
    def _get_align(attrs_dict: dict) -> str:
        style = attrs_dict.get("style", "")
        m = re.search(r"text-align:\s*(left|center|right|justify)", style, re.IGNORECASE)
        return m.group(1).lower() if m else "left"

    @staticmethod
    def _has_float(attrs_dict: dict) -> tuple[str | None, str | None]:
        """Return (float_direction, None) or (None, None)."""
        style = attrs_dict.get("style", "")
        m = re.search(r"float:\s*(left|right)", style, re.IGNORECASE)
        return (m.group(1).lower(), None) if m else (None, None)

    def _flush_buffer(self):
        text = "".join(self._buf).strip()
        # Collapse multiple spaces but keep single
        text = re.sub(r" {2,}", " ", text)
        self._buf.clear()
        return text

    # ── current context ───────────────────────────────────────────────────

    def _current_tag(self):
        return self._stack[-1] if self._stack else None

    # ── Parser events ─────────────────────────────────────────────────────

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        attrs_dict = dict(attrs)

        if tag in ("style", "script"):
            self._stack.append("__skip")
            return

        if tag == "br":
            self._buf.append("\n")
            return

        if tag in self.BLOCK_TAGS:
            # Flush any pending text into current block before opening new one
            # (in case of inline <br> inside a block)
            self._stack.append(tag)
            self._attrs_stack.append(attrs_dict)

        # For inline span with float — track separately
        if tag == "span":
            style = attrs_dict.get("style", "")
            fm = re.search(r"float:\s*(left|right)", style, re.IGNORECASE)
            if fm:
                direction = fm.group(1).lower()
                self._buf.append(f"\x00FLOAT_{direction.upper()}_START\x00")

    def handle_endtag(self, tag: str):
        tag = tag.lower()

        if tag in ("style", "script"):
            if self._stack and self._stack[-1] == "__skip":
                self._stack.pop()
            return

        if tag == "span":
            # Close any float span marker
            if "\x00FLOAT_" in "".join(self._buf):
                self._buf.append("\x00FLOAT_END\x00")
            return

        if tag in self.BLOCK_TAGS and self._stack and self._stack[-1] == tag:
            raw_text = self._flush_buffer()
            attrs_dict = self._attrs_stack.pop() if self._attrs_stack else {}
            self._stack.pop()

            if tag == "br":
                return

            # Skip empty blocks (except they're blank lines)
            if not raw_text and tag not in ("p",):
                return

            align = self._get_align(attrs_dict)

            # Check for float spans (date/city, signature/name)
            if "\x00FLOAT_LEFT_START\x00" in raw_text or "\x00FLOAT_RIGHT_START\x00" in raw_text:
                # Extract left and right parts
                left = right = ""
                # left part: between FLOAT_LEFT_START and FLOAT_END
                lm = re.search(r"\x00FLOAT_LEFT_START\x00(.*?)\x00FLOAT_END\x00", raw_text, re.DOTALL)
                rm = re.search(r"\x00FLOAT_RIGHT_START\x00(.*?)\x00FLOAT_END\x00", raw_text, re.DOTALL)
                if lm:
                    left = re.sub(r"\x00[^\\x00]*\x00", "", lm.group(1)).strip()
                if rm:
                    right = re.sub(r"\x00[^\\x00]*\x00", "", rm.group(1)).strip()
                if left or right:
                    self.items.append({"type": "tab_row", "left": left, "right": right})
                return

            # Strip float markers
            raw_text = re.sub(r"\x00[^\x00]*\x00", "", raw_text)
            if not raw_text and tag in ("div",):
                return

            if tag in ("h1", "h2", "h3"):
                self.items.append({"type": tag, "text": raw_text, "align": align})
            elif tag == "li":
                if raw_text:
                    self.items.append({"type": "li", "text": raw_text, "align": "justify"})
            else:
                if raw_text:
                    self.items.append({"type": "p", "text": raw_text, "align": align})
                else:
                    self.items.append({"type": "blank"})

    def handle_data(self, data: str):
        if self._stack and self._stack[-1] == "__skip":
            return
        self._buf.append(self._decode(data))

    def handle_entityref(self, name: str):
        entities = {"nbsp": " ", "amp": "&", "lt": "<", "gt": ">", "quot": '"'}
        self._buf.append(entities.get(name, ""))

    def handle_charref(self, name: str):
        if name.startswith("x"):
            c = chr(int(name[1:], 16))
        else:
            c = chr(int(name))
        self._buf.append(c if c != "\xa0" else " ")


def _parse_html_ordered(html: str) -> list[dict]:
    """Single-pass parse preserving document order."""
    parser = _DocParser()
    parser.feed(html)
    return parser.items


# ---------------------------------------------------------------------------
# DocumentService
# ---------------------------------------------------------------------------

class DocumentService:
    def __init__(self):
        settings = get_settings()
        self.generated_dir = Path(settings.storage_dir) / "generated"
        self.generated_dir.mkdir(parents=True, exist_ok=True)
        # A4 text width (210mm - 30mm left - 15mm right = 165mm) in twips (1cm = 567 twips)
        self._page_width_twips = int(16.5 * 567)

    # ── DOCX → TipTap HTML (template parsing) ─────────────────────────

    def docx_to_html(self, docx_path: str) -> str:
        """Parse a DOCX file and return structured HTML suitable for TipTap."""
        doc = Document(docx_path)
        html_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                html_parts.append("<p><br></p>")
                continue

            style_name = (para.style.name or "").lower()
            alignment = self._get_alignment(para)
            style_attr = f' style="text-align:{alignment}"' if alignment != "left" else ""

            if "heading 1" in style_name or style_name.startswith("heading 1"):
                html_parts.append(f"<h1{style_attr}>{self._runs_to_html(para)}</h1>")
            elif "heading 2" in style_name or style_name.startswith("heading 2"):
                html_parts.append(f"<h2{style_attr}>{self._runs_to_html(para)}</h2>")
            elif "heading 3" in style_name or style_name.startswith("heading 3"):
                html_parts.append(f"<h3{style_attr}>{self._runs_to_html(para)}</h3>")
            elif "list" in style_name:
                html_parts.append(f"<li>{self._runs_to_html(para)}</li>")
            else:
                centered_title = re.match(
                    r"^[ПП]\s*[РP]\s*[ЕE]\s*[ДD]\s*[СC]\s*[ТT]\s*[АA]\s*[ВB]\s*[ЛL]\s*[ЕE]\s*[НH]\s*[ИI]\s*[ЕE]",
                    text, re.IGNORECASE
                )
                if centered_title:
                    html_parts.append(
                        f'<h1 style="text-align:center"><strong>{text}</strong></h1>'
                    )
                else:
                    html_parts.append(f"<p{style_attr}>{self._runs_to_html(para)}</p>")

        html = "\n".join(html_parts)
        html = re.sub(r"(<li>.*?</li>(?:\s*<li>.*?</li>)*)", r"<ol>\1</ol>", html, flags=re.DOTALL)
        return html

    def _get_alignment(self, para) -> str:
        alignment = para.alignment
        if alignment is None and para.style and para.style.paragraph_format:
            alignment = para.style.paragraph_format.alignment
        return ALIGN_MAP.get(alignment, "left")

    def _runs_to_html(self, para) -> str:
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            parts.append(text)
        return "".join(parts) if parts else para.text

    # ── Generate structured article-200 template HTML ──────────────────

    @staticmethod
    def representation_template_html() -> str:
        return """\
<p style="text-align:right"><strong>Руководителю ___________________________</strong></p>
<p style="text-align:right">________________________________</p>
<p style="text-align:right">________________________________</p>
<p><br></p>
<h1 style="text-align:center"><strong>П Р Е Д С Т А В Л Е Н И Е</strong></h1>
<h2 style="text-align:center">по устранению обстоятельств, способствовавших совершению уголовного правонарушения</h2>
<p><br></p>
<p><span style="float:left">«___» _____________ 20___ г.</span><span style="float:right">г. _______________</span></p>
<div style="clear:both"></div>
<p><br></p>
<p style="text-align:justify">Следственным отделом ___________________________ расследуется уголовное дело, зарегистрированное в ЕРДР за № _______________, возбужденное по признакам уголовного правонарушения, предусмотренного _____ Уголовного кодекса Республики Казахстан.</p>
<p><br></p>
<p style="text-align:justify">В ходе расследования данного уголовного дела установлено, что совершению указанного уголовного правонарушения способствовали следующие обстоятельства:</p>
<p><br></p>
<p style="text-align:justify">___________________________________________________________________________</p>
<p><br></p>
<p style="text-align:justify">На основании изложенного, руководствуясь статьей 200 Уголовно-процессуального кодекса Республики Казахстан,</p>
<p><br></p>
<h2 style="text-align:center"><strong>ПРЕДЛАГАЮ:</strong></h2>
<p><br></p>
<ol>
<li><p style="text-align:justify">Принять необходимые меры по устранению выявленных обстоятельств, способствовавших совершению уголовного правонарушения.</p></li>
<li><p style="text-align:justify">О принятых мерах сообщить в следственный отдел в месячный срок со дня получения настоящего представления.</p></li>
</ol>
<p><br></p>
<p>Следователь ________________________________</p>
<p><span style="float:left">_________________ (звание)</span><span style="float:right">________________ (Ф.И.О.)</span></p>
<div style="clear:both"></div>
"""

    # ── HTML → DOCX export ─────────────────────────────────────────────

    def html_to_docx(self, html_content: str, filename: str = "document.docx") -> str:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            self._page_width_twips = int(
                (section.page_width - section.left_margin - section.right_margin) / 914400 * 1440
            )

        # Set Normal style
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Cm(1.25)

        items = _parse_html_ordered(html_content)

        for item in items:
            itype = item["type"]

            if itype == "blank":
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)

            elif itype == "h1":
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = self._html_align_to_wd(item.get("align", "center"))
                run = p.add_run(item["text"])
                run.font.name = "Times New Roman"
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif itype in ("h2", "h3"):
                # Do NOT use Word heading style (it's blue) — use regular paragraph
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = self._html_align_to_wd(item.get("align", "center"))
                run = p.add_run(item["text"])
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif itype == "li":
                p = doc.add_paragraph(style="List Number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(item["text"])
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif itype == "tab_row":
                self._add_tab_paragraph(doc, item["left"], item["right"])

            elif itype == "p":
                text = item["text"]
                align = self._html_align_to_wd(item.get("align", "left"))
                p = doc.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_after = Pt(0)
                if align in (WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER):
                    p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

        safe_filename = f"{uuid.uuid4().hex}_{filename}"
        file_path = self.generated_dir / safe_filename
        doc.save(str(file_path))
        return str(file_path)

    def _add_tab_paragraph(self, doc: Document, left_text: str, right_text: str):
        """Add paragraph: left_text [TAB→right edge] right_text."""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)

        # Right-aligned tab-stop at text-area right edge
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:pos"), str(int(self._page_width_twips)))
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

        if left_text:
            rl = p.add_run(left_text)
            rl.font.name = "Times New Roman"
            rl.font.size = Pt(14)
            rl.font.color.rgb = RGBColor(0, 0, 0)

        # TAB character element
        tab_run = p.add_run()
        tab_run.font.name = "Times New Roman"
        tab_run.font.size = Pt(14)
        tab_run._r.append(OxmlElement("w:tab"))

        if right_text:
            rr = p.add_run(right_text)
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(14)
            rr.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _html_align_to_wd(align: str):
        return {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    # Keep old static for backward compat
    @staticmethod
    def _parse_html(html: str) -> list[dict]:
        return _parse_html_ordered(html)


def _strip_tags(html: str) -> str:
    return re.sub(r"<[^>]+>", "", html).strip()
