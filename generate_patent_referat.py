# -*- coding: utf-8 -*-
"""
Генерация заполненного Реферата на программу для ЭВМ «ПредставлениеAi»
для регистрации авторского права в соответствии со статьёй 9-1
Закона РК «Об авторском праве и смежных правах» от 10 июня 1996 года № 6.

Соответствует требованиям KazPatent (Национального института интеллектуальной
собственности) для подачи заявки на регистрацию программы для ЭВМ.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

def set_run_font(run, name='Times New Roman', size=14, bold=False, italic=False, underline=False):
    """Утилита для установки шрифта run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    # Для корректного отображения кириллицы
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)

def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, 
                  italic=False, underline=False, size=14, space_after=0, space_before=0):
    """Добавить параграф с заданным форматированием."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, underline=underline)
    return p

def add_section_header(doc, text):
    """Добавить заголовок раздела (жирный, подчёркнутый, по центру)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, underline=True)
    return p

def add_section_body(doc, text):
    """Добавить тело раздела (обычный текст, выравнивание по ширине)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run, size=14)
    return p

def generate_referat():
    """Генерация заполненного реферата."""
    doc = Document()
    
    # ── Настройка полей страницы ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # ═══════════════════════════════════════════════════
    #  ШАПКА (правый верхний угол)
    # ═══════════════════════════════════════════════════
    
    header_lines = [
        'Статья 9-1 Закона РК',
        '«Об авторском праве',
        'и смежных правах»',
        'от 10 июня 1996 года № 6'
    ]
    for line in header_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, size=12, bold=False, italic=True)
    
    # ── Пустые строки ──
    for _ in range(3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    
    # ═══════════════════════════════════════════════════
    #  ЗАГОЛОВОК: «Реферат»
    # ═══════════════════════════════════════════════════
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Реферат')
    set_run_font(run, size=16, bold=True)
    
    # ── Вид объекта ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('программа для ЭВМ')
    set_run_font(run, size=14, bold=True)
    
    # ── Название ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('под названием')
    set_run_font(run, size=14, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('«ПредставлениеAi — Платформа юридического помощника на базе '
                     'искусственного интеллекта для Министерства внутренних дел '
                     'Республики Казахстан»')
    set_run_font(run, size=14, bold=True)
    
    # ═══════════════════════════════════════════════════
    #  АВТОР(Ы)
    # ═══════════════════════════════════════════════════
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    
    run = p.add_run('Фамилия, имя, отчество автора(-ов): ')
    set_run_font(run, size=14, bold=True, underline=True)
    run2 = p.add_run('______________________________________')
    set_run_font(run2, size=14)
    
    # ── Дата создания ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    run = p.add_run('Дата создания объекта: ')
    set_run_font(run, size=14, bold=True, underline=True)
    run2 = p.add_run('«___» ____________ 2025 г.')
    set_run_font(run2, size=14)
    
    # ═══════════════════════════════════════════════════
    #  1. ОБЛАСТЬ ПРИМЕНЕНИЯ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Область применения')
    
    add_section_body(doc,
        'Программа для ЭВМ «ПредставлениеAi» предназначена для применения '
        'в сфере правоохранительной деятельности Республики Казахстан, '
        'а именно — в органах следствия и дознания Министерства внутренних '
        'дел Республики Казахстан (далее — МВД РК).'
    )
    add_section_body(doc,
        'Область применения программы охватывает автоматизацию '
        'юридической деятельности следователей и дознавателей при работе с '
        'уголовными делами, включая: анализ материалов уголовных дел с '
        'применением технологий искусственного интеллекта (далее — ИИ); '
        'контекстный поиск по загруженным процессуальным документам; '
        'автоматическую генерацию процессуальных документов (представлений) '
        'в соответствии со статьёй 200 Уголовно-процессуального кодекса '
        'Республики Казахстан; интеллектуальный поиск и цитирование '
        'релевантных норм законодательства посредством технологии RAG '
        '(Retrieval-Augmented Generation).'
    )
    add_section_body(doc,
        'Программа может быть использована государственными органами, '
        'осуществляющими уголовное преследование, юридическими '
        'подразделениями правоохранительных органов, а также '
        'образовательными учреждениями для подготовки кадров в области '
        'юриспруденции и правоприменения.'
    )
    
    # ═══════════════════════════════════════════════════
    #  2. НАЗНАЧЕНИЕ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Назначение')
    
    add_section_body(doc,
        'Программа «ПредставлениеAi» предназначена для решения следующих задач:'
    )
    
    tasks = [
        'автоматизация процесса создания процессуальных документов — '
        'представлений в порядке статьи 200 УПК РК, с обеспечением '
        'строгого соответствия требованиям уголовно-процессуального '
        'законодательства;',
        
        'интеллектуальный анализ материалов уголовных дел, загруженных '
        'в систему в формате PDF, DOCX, TXT и изображений, с '
        'применением векторного поиска и языковых моделей ИИ;',
        
        'контекстный диалоговый поиск (RAG-чат) по материалам '
        'уголовного дела, позволяющий следователю получать ответы на '
        'вопросы непосредственно по содержимому загруженных документов;',
        
        'автоматический поиск и подбор релевантных норм '
        'законодательства Республики Казахстан с использованием '
        'гибридной системы поиска (FAISS + BM25), обеспечивающей '
        'высокую точность результатов;',
        
        'генерация юридически грамотных текстов процессуальных '
        'документов с возможностью редактирования в встроенном '
        'WYSIWYG-редакторе и последующего экспорта в формат DOCX '
        'по стандартам ГОСТ (Times New Roman, 14pt, междустрочный '
        'интервал 1.5, поля по ГОСТ);',
        
        'управление уголовными делами, включая регистрацию по номеру '
        'ЕРДР, хранение фабулы, загрузку и индексацию процессуальных '
        'документов.'
    ]
    
    for i, task in enumerate(tasks, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.25)
        run = p.add_run(f'{i}) {task}')
        set_run_font(run, size=14)
    
    # ═══════════════════════════════════════════════════
    #  3. ФУНКЦИОНАЛЬНЫЕ ВОЗМОЖНОСТИ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Функциональные возможности')
    
    add_section_body(doc,
        'Программа «ПредставлениеAi» обеспечивает следующие '
        'функциональные возможности:'
    )
    
    features = [
        ('Модуль управления уголовными делами', 
         'регистрация дел по номеру ЕРДР (Единый реестр досудебных '
         'расследований) с указанием фабулы дела; хранение метаданных '
         'уголовного дела; привязка загруженных документов к '
         'конкретному делу.'),
        
        ('Модуль загрузки и обработки документов',
         'загрузка файлов в форматах PDF, DOCX, TXT и изображений '
         'объёмом до 50 МБ; автоматическое извлечение текста из '
         'загруженных документов; векторная индексация содержимого '
         'документов для семантического поиска с использованием модели '
         'embeddings intfloat/multilingual-e5-base и хранилища '
         'ChromaDB.'),
        
        ('Модуль RAG-чата (Retrieval-Augmented Generation)',
         'диалоговый интерфейс для вопросно-ответного '
         'взаимодействия с ИИ по материалам уголовного дела; '
         'контекстный поиск по загруженным документам с '
         'предоставлением источников; использование языковой '
         'модели Gemini 2.5 Flash через OpenAI-совместимый API.'),
        
        ('Модуль генерации процессуальных документов',
         'пошаговый мастер (wizard) для формирования представлений '
         'по статье 200 УПК РК; выбор типа нарушения, адресата '
         'представления, сроков исполнения; автоматический поиск '
         'и подстановка релевантных норм законодательства; '
         'генерация юридически грамотного текста с соблюдением '
         'структуры процессуального документа.'),
        
        ('Модуль поиска по законодательству (RAG Laws)',
         'гибридная система поиска по обогащённой базе '
         'законодательства РК (FAISS + BM25); индексация '
         'более 200 статей УПК РК с обогащёнными метаданными; '
         'мульти-запросный поиск для повышения полноты и '
         'точности результатов.'),
        
        ('WYSIWYG-редактор документов',
         'встроенный визуальный редактор на базе TipTap с '
         'поддержкой форматирования шрифтов, размеров, ссылок, '
         'цитат, выделения текста и выравнивания; экспорт '
         'отредактированных документов в формат DOCX по '
         'стандартам ГОСТ.'),
        
        ('Модуль аутентификации и ролевого доступа',
         'регистрация и авторизация пользователей с '
         'разграничением ролей: «Администратор» (управление '
         'шаблонами, настройками системы) и «Следователь» '
         '(работа с делами, документами, чатом); первый '
         'зарегистрированный пользователь автоматически '
         'получает роль администратора.'),
        
        ('Модуль пользовательских шаблонов',
         'загрузка и управление шаблонами документов в '
         'форматах DOCX, RTF, ODT через панель администратора; '
         'генерация документов по загруженным шаблонам с '
         'автоматическим заполнением фактами из уголовного дела.')
    ]
    
    for i, (title, desc) in enumerate(features, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        
        run_num = p.add_run(f'{i}. ')
        set_run_font(run_num, size=14, bold=True)
        
        run_title = p.add_run(f'{title} — ')
        set_run_font(run_title, size=14, bold=True)
        
        run_desc = p.add_run(desc)
        set_run_font(run_desc, size=14)
    
    # ═══════════════════════════════════════════════════
    #  4. ОСНОВНЫЕ ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Основные технические характеристики')
    
    add_section_body(doc,
        'Программа «ПредставлениеAi» построена на основе '
        'клиент-серверной архитектуры и развёртывается с помощью '
        'технологии контейнеризации Docker. Система состоит из '
        'следующих компонентов:'
    )
    
    tech_specs = [
        ('Серверная часть (Backend)', [
            'Веб-фреймворк: FastAPI (Python 3.12)',
            'ORM: SQLAlchemy (асинхронный режим)',
            'Миграции базы данных: Alembic',
            'API-сервер: Uvicorn (ASGI)',
            'Протокол: REST API',
            'Версия: 2.0.0',
        ]),
        ('Клиентская часть (Frontend)', [
            'Фреймворк: Next.js 15 (React 19)',
            'Язык: TypeScript',
            'Стилизация: Tailwind CSS',
            'Визуальный редактор: TipTap (WYSIWYG)',
        ]),
        ('База данных', [
            'СУБД: PostgreSQL 16 (Alpine)',
            'Векторное хранилище: ChromaDB (для семантического поиска по документам)',
            'Индекс: FAISS (Facebook AI Similarity Search) для поиска по законодательству',
        ]),
        ('Модели искусственного интеллекта', [
            'Языковая модель: Gemini 2.5 Flash (через OpenAI-совместимый API, Vertex AI)',
            'Модель embeddings: intfloat/multilingual-e5-base (мультиязычные векторные представления)',
            'Методология: RAG (Retrieval-Augmented Generation) — генерация с дополненным извлечением',
        ]),
        ('Инфраструктура', [
            'Контейнеризация: Docker, Docker Compose',
            'Максимальный размер загружаемого файла: 50 МБ (52 428 800 байт)',
            'Поддерживаемые форматы ввода: PDF, DOCX, TXT, изображения (JPEG, PNG)',
            'Формат вывода документов: DOCX (по стандартам ГОСТ)',
            'Протокол взаимодействия с ИИ: OpenAI-совместимый API',
        ]),
    ]
    
    for group_title, items in tech_specs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(group_title + ':')
        set_run_font(run, size=14, bold=True)
        
        for item in items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.left_indent = Cm(1.25)
            
            # Разделяем на название и значение по двоеточию
            if ': ' in item:
                key, val = item.split(': ', 1)
                run_key = p.add_run(f'— {key}: ')
                set_run_font(run_key, size=14, bold=False)
                run_val = p.add_run(val)
                set_run_font(run_val, size=14)
            else:
                run = p.add_run(f'— {item}')
                set_run_font(run, size=14)
    
    # ═══════════════════════════════════════════════════
    #  5. ЯЗЫК ПРОГРАММИРОВАНИЯ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Язык программирования')
    
    add_section_body(doc,
        'Программа «ПредставлениеAi» разработана с использованием '
        'следующих языков программирования и технологий:'
    )
    
    langs = [
        ('Python 3.12', 'основной язык серверной части (backend). '
         'Используется для реализации REST API, бизнес-логики, '
         'взаимодействия с базой данных, моделями ИИ, '
         'векторных хранилищ и системы RAG.'),
        
        ('TypeScript', 'основной язык клиентской части (frontend). '
         'Используется в связке с фреймворком Next.js 15 / React 19 '
         'для построения пользовательского интерфейса.'),
        
        ('SQL', 'язык запросов к реляционной базе данных PostgreSQL 16 '
         'через ORM SQLAlchemy.'),
        
        ('HTML5/CSS3', 'языки разметки и стилизации '
         'пользовательского интерфейса (Tailwind CSS).'),
        
        ('YAML', 'язык описания конфигурации контейнеров '
         '(Docker Compose) и CI/CD процессов.'),
    ]
    
    for lang, desc in langs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        
        run_lang = p.add_run(f'{lang} — ')
        set_run_font(run_lang, size=14, bold=True)
        
        run_desc = p.add_run(desc)
        set_run_font(run_desc, size=14)
    
    # ═══════════════════════════════════════════════════
    #  6. ТИП РЕАЛИЗУЮЩЕЙ ЭВМ
    # ═══════════════════════════════════════════════════
    
    add_section_header(doc, 'Тип реализующей ЭВМ')
    
    add_section_body(doc,
        'Программа «ПредставлениеAi» является веб-приложением, '
        'функционирующим в архитектуре «клиент-сервер» и '
        'предназначенным для работы на следующих типах ЭВМ:'
    )
    
    add_section_body(doc,
        'Серверная часть: любой серверный компьютер (физический '
        'или виртуальный) под управлением ОС Linux (рекомендуется '
        'Ubuntu 22.04 LTS или аналогичный дистрибутив) с '
        'установленными Docker и Docker Compose. Минимальные '
        'требования: процессор x86_64, оперативная память не менее '
        '8 ГБ, дисковое пространство не менее 20 ГБ. Для работы '
        'моделей ИИ рекомендуется наличие GPU NVIDIA с поддержкой '
        'CUDA (опционально).'
    )
    
    add_section_body(doc,
        'Клиентская часть: любой персональный компьютер, ноутбук '
        'или мобильное устройство с современным веб-браузером '
        '(Google Chrome версии 90 и выше, Mozilla Firefox версии '
        '88 и выше, Microsoft Edge версии 90 и выше, Safari версии '
        '14 и выше). Доступ к программе осуществляется через '
        'веб-интерфейс по протоколу HTTP/HTTPS без необходимости '
        'установки дополнительного программного обеспечения на '
        'стороне клиента.'
    )
    
    add_section_body(doc,
        'Программа поддерживает одновременную работу нескольких '
        'пользователей благодаря асинхронной архитектуре серверной '
        'части (FastAPI/Uvicorn ASGI) и может быть развёрнута как '
        'в локальной сети организации, так и в облачной '
        'инфраструктуре.'
    )
    
    # ═══════════════════════════════════════════════════
    #  Сохранение
    # ═══════════════════════════════════════════════════
    
    output_path = 'Referat_PredstavlenieAi.docx'
    doc.save(output_path)
    print(f'Referat generated: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_referat()
